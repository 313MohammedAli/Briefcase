"""PDF (WeasyPrint) and DOCX (python-docx) export for cover letters and resumes."""

import io
from html import escape

PDF_STYLE = """
  @page { size: letter; margin: 1in; }
  body { font-family: Helvetica, Arial, sans-serif; font-size: 11pt; color: #1a1a1a; line-height: 1.5; }
  h1 { font-size: 16pt; margin: 0 0 2pt 0; }
  h2 { font-size: 12pt; border-bottom: 1px solid #999; padding-bottom: 2pt; margin: 14pt 0 6pt 0; }
  .subtitle { color: #555; margin: 0 0 14pt 0; }
  .entry-header { font-weight: bold; margin: 8pt 0 2pt 0; }
  .entry-meta { color: #555; font-size: 10pt; margin: 0 0 4pt 0; }
  ul { margin: 2pt 0 8pt 0; padding-left: 16pt; }
  li { margin-bottom: 2pt; }
  p { margin: 0 0 10pt 0; }
"""


def _pdf_from_body(body_html: str) -> bytes:
    # Imported lazily: WeasyPrint needs system libs (pango/gobject) that may
    # be absent in local dev; DOCX export and the rest of the API still work.
    from weasyprint import HTML

    html = f"<html><head><style>{PDF_STYLE}</style></head><body>{body_html}</body></html>"
    return HTML(string=html).write_pdf()


def cover_letter_pdf(job_application, paragraphs: list[str]) -> bytes:
    body = (
        f"<h1>{escape(job_application.job_title)}</h1>"
        f"<p class='subtitle'>{escape(job_application.company)}</p>"
        + "".join(f"<p>{escape(p)}</p>" for p in paragraphs)
    )
    return _pdf_from_body(body)


def cover_letter_docx(job_application, paragraphs: list[str]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(job_application.job_title, level=1)
    doc.add_paragraph(job_application.company)
    for p in paragraphs:
        doc.add_paragraph(p)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def resume_pdf(job_application, resume: dict, candidate_name: str) -> bytes:
    parts = [f"<h1>{escape(candidate_name)}</h1>" if candidate_name else ""]
    parts.append(
        f"<p class='subtitle'>Tailored for {escape(job_application.job_title)} at "
        f"{escape(job_application.company)}</p>"
    )
    if resume.get("summary"):
        parts.append(f"<h2>Summary</h2><p>{escape(resume['summary'])}</p>")
    parts.append("<h2>Experience</h2>")
    for entry in resume.get("entries", []):
        header = escape(entry.get("title", ""))
        if entry.get("organization"):
            header += f" — {escape(entry['organization'])}"
        parts.append(f"<p class='entry-header'>{header}</p>")
        if entry.get("dates"):
            parts.append(f"<p class='entry-meta'>{escape(entry['dates'])}</p>")
        bullets = "".join(f"<li>{escape(b)}</li>" for b in entry.get("bullets", []))
        parts.append(f"<ul>{bullets}</ul>")
    return _pdf_from_body("".join(parts))


def resume_docx(job_application, resume: dict, candidate_name: str) -> bytes:
    from docx import Document

    doc = Document()
    if candidate_name:
        doc.add_heading(candidate_name, level=1)
    doc.add_paragraph(
        f"Tailored for {job_application.job_title} at {job_application.company}"
    )
    if resume.get("summary"):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(resume["summary"])
    doc.add_heading("Experience", level=2)
    for entry in resume.get("entries", []):
        header = entry.get("title", "")
        if entry.get("organization"):
            header += f" — {entry['organization']}"
        doc.add_heading(header, level=3)
        if entry.get("dates"):
            doc.add_paragraph(entry["dates"])
        for bullet in entry.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
