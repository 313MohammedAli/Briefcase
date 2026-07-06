"""Resume upload: text extraction and LLM parsing into experience entries.

The uploaded file is reduced to plain text (PDF via pypdf, DOCX via
python-docx, or raw text), then an LLM structures it into proposed
ExperienceEntry payloads that the user reviews and edits in the UI
before anything is saved to their bank.
"""

import io

from rest_framework.exceptions import ValidationError

from .generation import _structured_call

MAX_UPLOAD_BYTES = 5 * 1024 * 1024

ENTRY_TYPES = ["job", "project", "certification", "education"]


def extract_text_from_upload(uploaded_file) -> str:
    """Returns the plain text of an uploaded PDF, DOCX, or text resume."""
    if uploaded_file.size > MAX_UPLOAD_BYTES:
        raise ValidationError("Resume file is too large (5 MB max).")

    name = (uploaded_file.name or "").lower()
    data = uploaded_file.read()

    if name.endswith(".pdf"):
        from pypdf import PdfReader

        try:
            reader = PdfReader(io.BytesIO(data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            raise ValidationError("Could not read this PDF. Try exporting it again or use DOCX.")
    elif name.endswith(".docx"):
        from docx import Document

        try:
            document = Document(io.BytesIO(data))
        except Exception:
            raise ValidationError("Could not read this DOCX file.")
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        text = "\n".join(parts)
    elif name.endswith((".txt", ".md")):
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1")
    else:
        raise ValidationError("Unsupported file type. Upload a PDF, DOCX, or TXT resume.")

    text = text.strip()
    if len(text) < 50:
        raise ValidationError(
            "Couldn't extract readable text from this file. If it's a scanned PDF, "
            "export a text-based version and try again."
        )
    return text


def extract_resume_entries(resume_text: str) -> list[dict]:
    """LLM-structures resume text into proposed experience entries."""
    entry_schema = {
        "type": "object",
        "properties": {
            "type": {"type": "string", "enum": ENTRY_TYPES},
            "title": {"type": "string"},
            "organization": {"type": "string"},
            "start_date": {
                "type": "string",
                "description": "ISO date YYYY-MM-DD, first of month if only month known, empty if unknown",
            },
            "end_date": {
                "type": "string",
                "description": "ISO date YYYY-MM-DD, empty if unknown or current/ongoing",
            },
            "tags": {"type": "array", "items": {"type": "string"}},
            "bullets": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["type", "title", "organization", "start_date", "end_date", "tags", "bullets"],
        "additionalProperties": False,
    }
    schema = {
        "type": "object",
        "properties": {"entries": {"type": "array", "items": entry_schema}},
        "required": ["entries"],
        "additionalProperties": False,
    }
    system = (
        "You convert resume text into structured experience entries. Create one "
        "entry per job, project, certification, or education item. Copy bullet "
        "points faithfully from the resume (light cleanup only — fix broken line "
        "wraps, keep the author's wording). Put skills/technologies mentioned for "
        "that entry into tags. A standalone skills section should be distributed "
        "into the most relevant entries' tags, not invented as its own entry. "
        "Dates: ISO YYYY-MM-DD, use the first of the month when only month/year "
        "is given, empty string when unknown or ongoing. Do not invent anything "
        "that is not in the resume."
    )
    result = _structured_call(system, resume_text, "resume_entries", schema)
    entries = result["entries"]
    for entry in entries:
        # Normalize empty-string dates to None to match the serializer contract.
        entry["start_date"] = entry["start_date"] or None
        entry["end_date"] = entry["end_date"] or None
    return entries
