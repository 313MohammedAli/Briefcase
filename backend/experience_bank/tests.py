from django.test import TestCase
from rest_framework.test import APIClient

from accounts.models import Profile
from common.generation import keyword_gap_analysis
from common.retrieval import _score_from_distances, retrieve_relevant_bullets

from .models import ExperienceBullet, ExperienceEntry


def make_entry(owner, title="Engineer", bullets=(), tags=()):
    entry = ExperienceEntry.objects.create(
        owner=owner, type="job", title=title, organization="Acme", tags=list(tags)
    )
    for i, text in enumerate(bullets):
        ExperienceBullet.objects.create(entry=entry, text=text, order=i)
    return entry


class RetrievalTests(TestCase):
    def setUp(self):
        self.owner = Profile.objects.create(clerk_user_id="user_1", email="a@b.c")

    def test_small_bank_returns_all_bullets_without_retrieval(self):
        make_entry(self.owner, bullets=["Built an API", "Led a team"])
        result = retrieve_relevant_bullets(self.owner, [0.1] * 1536)
        self.assertFalse(result.used_retrieval)
        self.assertEqual(len(result.bullets), 2)

    def test_no_embedding_returns_all_bullets_unscored(self):
        make_entry(self.owner, bullets=["Built an API"])
        result = retrieve_relevant_bullets(self.owner, None)
        self.assertFalse(result.used_retrieval)
        self.assertIsNone(result.match_score)

    def test_score_mapping_clamps_to_0_100(self):
        self.assertEqual(_score_from_distances([0.9]), 0.0)  # similarity 0.1 -> floor
        self.assertEqual(_score_from_distances([0.4]), 100.0)  # similarity 0.6 -> ceiling
        self.assertEqual(_score_from_distances([0.65]), 50.0)  # similarity 0.35 -> midpoint


class KeywordGapTests(TestCase):
    def setUp(self):
        self.owner = Profile.objects.create(clerk_user_id="user_2", email="b@c.d")
        self.entry = make_entry(
            self.owner,
            bullets=["Built Django REST APIs on PostgreSQL"],
            tags=["Python"],
        )

    def test_literal_and_tag_matches_without_embeddings(self):
        bullets = list(ExperienceBullet.objects.filter(entry__owner=self.owner))
        result = keyword_gap_analysis(["Django", "Python", "Kubernetes"], bullets)
        self.assertEqual(result["matched"], ["Django", "Python"])
        self.assertEqual(result["missing"], ["Kubernetes"])

    def test_empty_keywords(self):
        self.assertEqual(keyword_gap_analysis([], []), {"matched": [], "missing": []})


class PerUserScopingTests(TestCase):
    def setUp(self):
        self.alice = Profile.objects.create(clerk_user_id="user_a", email="alice@x.y")
        self.bob = Profile.objects.create(clerk_user_id="user_b", email="bob@x.y")
        self.alice_entry = make_entry(self.alice, title="Alice's job", bullets=["did things"])
        self.client = APIClient()

    def test_user_only_sees_own_entries(self):
        self.client.force_authenticate(user=self.bob)
        response = self.client.get("/api/experience-entries/")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json(), [])

    def test_user_cannot_fetch_others_entry(self):
        self.client.force_authenticate(user=self.bob)
        response = self.client.get(f"/api/experience-entries/{self.alice_entry.id}/")
        self.assertEqual(response.status_code, 404)

    def test_create_and_update_entry_with_bullets(self):
        self.client.force_authenticate(user=self.bob)
        response = self.client.post(
            "/api/experience-entries/",
            {
                "type": "project",
                "title": "Side project",
                "tags": ["React"],
                "bullets": [{"text": "Shipped it", "order": 0}],
            },
            format="json",
        )
        self.assertEqual(response.status_code, 201)
        entry_id = response.json()["id"]

        response = self.client.put(
            f"/api/experience-entries/{entry_id}/",
            {
                "type": "project",
                "title": "Side project v2",
                "tags": [],
                "bullets": [{"text": "Rewrote it", "order": 0}, {"text": "Scaled it", "order": 1}],
            },
            format="json",
        )
        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(body["title"], "Side project v2")
        self.assertEqual(len(body["bullets"]), 2)


class ResumeImportTests(TestCase):
    def setUp(self):
        self.owner = Profile.objects.create(clerk_user_id="user_r", email="r@x.y")
        self.client = APIClient()
        self.client.force_authenticate(user=self.owner)

    def _upload(self, name, content):
        from django.core.files.uploadedfile import SimpleUploadedFile

        return self.client.post(
            "/api/experience-entries/import-resume/",
            {"file": SimpleUploadedFile(name, content)},
            format="multipart",
        )

    def test_requires_api_key(self):
        with self.settings(OPENAI_API_KEY=""):
            response = self._upload("resume.txt", b"x" * 100)
        self.assertEqual(response.status_code, 503)

    def test_rejects_unsupported_file_type(self):
        with self.settings(OPENAI_API_KEY="test-key"):
            response = self._upload("resume.png", b"x" * 100)
        self.assertEqual(response.status_code, 400)

    def test_rejects_unreadable_content(self):
        with self.settings(OPENAI_API_KEY="test-key"):
            response = self._upload("resume.txt", b"hi")
        self.assertEqual(response.status_code, 400)

    def test_extracts_docx_text(self):
        import io

        from docx import Document

        from common.resume_import import extract_text_from_upload

        doc = Document()
        doc.add_paragraph("Software Engineer at Acme Corp, 2020 to present.")
        doc.add_paragraph("Built Django services handling millions of requests.")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        buffer.name = "resume.docx"
        buffer.size = len(buffer.getvalue())
        text = extract_text_from_upload(buffer)
        self.assertIn("Acme Corp", text)
        self.assertIn("Django services", text)

    def test_bulk_create(self):
        response = self.client.post(
            "/api/experience-entries/bulk-create/",
            {
                "entries": [
                    {
                        "type": "job",
                        "title": "Engineer",
                        "organization": "Acme",
                        "tags": ["Python"],
                        "bullets": [{"text": "Did things", "order": 0}],
                    },
                    {
                        "type": "education",
                        "title": "BSc Computer Science",
                        "organization": "State University",
                        "tags": [],
                        "bullets": [],
                    },
                ]
            },
            format="json",
        )
        self.assertEqual(response.status_code, 201)
        self.assertEqual(ExperienceEntry.objects.filter(owner=self.owner).count(), 2)

    def test_bulk_create_rejects_empty(self):
        response = self.client.post(
            "/api/experience-entries/bulk-create/", {"entries": []}, format="json"
        )
        self.assertEqual(response.status_code, 400)


class HumanStyleTests(TestCase):
    def test_strip_dashes_replaces_em_and_en_dashes(self):
        from common.generation import _strip_dashes

        self.assertEqual(
            _strip_dashes("I led the migration — a six month effort — to Django."),
            "I led the migration, a six month effort, to Django.",
        )
        self.assertEqual(_strip_dashes("2019–2022 at Acme"), "2019, 2022 at Acme")

    def test_strip_dashes_keeps_word_hyphens(self):
        from common.generation import _strip_dashes

        self.assertEqual(
            _strip_dashes("Cross-functional, self-starter role"),
            "Cross-functional, self-starter role",
        )


class UploadGuardTests(TestCase):
    def test_docx_zip_bomb_rejected(self):
        import io
        import zipfile

        from rest_framework.exceptions import ValidationError

        from common.resume_import import extract_text_from_upload

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("word/document.xml", b"A" * (60 * 1024 * 1024))
        buf.seek(0)
        buf.name = "resume.docx"
        buf.size = len(buf.getvalue())
        with self.assertRaises(ValidationError):
            extract_text_from_upload(buf)
